#!/usr/bin/env python3
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _add_runs(paragraph, text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _parse_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def _is_table_separator(line):
    cells = _parse_table_row(line)
    if not cells:
        return False
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    _set_doc_font(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    i = 0
    in_code = False
    code_lang = ""
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            else:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                in_code = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            headers = _parse_table_row(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for col, header in enumerate(headers):
                cell = table.rows[0].cells[col]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(header)
                run.bold = True
            for r_idx, row in enumerate(rows):
                for c_idx in range(len(headers)):
                    value = row[c_idx] if c_idx < len(row) else ""
                    table.rows[r_idx + 1].cells[c_idx].text = value
            doc.add_paragraph()
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, numbered.group(2))
            i += 1
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line.strip()[2:])
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.italic = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs(p, line.strip())
        i += 1

    doc.save(str(docx_path))


def main():
    if len(sys.argv) != 3:
        print("Usage: md_to_docx.py input.md output.docx")
        sys.exit(1)
    convert_md_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"Written: {sys.argv[2]}")


if __name__ == "__main__":
    main()
